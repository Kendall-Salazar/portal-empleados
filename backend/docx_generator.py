"""
Generador de documentos Word (.docx) para acciones de empleados.
Diseño corporativo profesional — Servicentro La Marina.

Tipos de documento:
  1. Contrato de Préstamo (+ tabla de pagos en hoja aparte)
  2. Carta de Amonestación (faltantes / tardanzas / conductas)
  3. Constancia de Vacaciones
  4. Carta de Despido
  5. Carta de Aceptación de Renuncia
  6. Carta de Recomendación
  7. Liquidación Laboral
"""

import os
import math
import subprocess
import json
import pathlib
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

EMPRESA_NOMBRE = "Servicentro La Marina"
EMPRESA_CEDULA = "Cédula Jurídica 3-101-130178"
EMPRESA_TEL = "Tel: 2479-3131"
EMPRESA_DIR = "San Carlos, Alajuela, Costa Rica"

# ── Reference counter storage ─────────────────────────────────────────────────
DATA_FILE = pathlib.Path(__file__).parent / "reference_counter.json"

TIPO_CODIGO = {
    "prestamo": "PR",
    "amonestacion": "AMON",
    "despido": "DESP",
    "renuncia": "REN",
    "liquidacion": "LIQ",
    "liquidacion_laboral": "LIQ",
    "recomendacion": "CR",
    "constancia_salarial": "CS",
}

def get_next_reference(tipo: str) -> str:
    """Genera el siguiente código de referencia para el tipo de documento."""
    year = datetime.now().year
    prefix = TIPO_CODIGO.get(tipo.lower(), "DOC")
    if DATA_FILE.exists():
        data = json.loads(DATA_FILE.read_text(encoding="utf-8"))
    else:
        data = {}
    key = f"{prefix}-{year}"
    current_num = data.get(key, 0) + 1
    data[key] = current_num
    DATA_FILE.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    return f"{prefix}-{year}-{current_num:04d}"


# ── Color constants ───────────────────────────────────────────────────────────
BLUE_DARK = RGBColor(0x1A, 0x3C, 0x6E)
BLUE_MED  = RGBColor(0x2F, 0x54, 0x96)
BLUE_LIGHT_HEX = 'E8EEF7'
GRAY_TEXT = RGBColor(0x64, 0x74, 0x8B)
GRAY_LINE = RGBColor(0xBD, 0xC5, 0xD6)
BLACK = RGBColor(0x0F, 0x17, 0x2A)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ─── Helpers ──────────────────────────────────────────────────────────────────

def _build_output_path(base_dir: str, tipo_doc: str, emp_nombre: str) -> str:
    today = datetime.now()
    month_folder = today.strftime("%Y-%m")
    safe_name = emp_nombre.replace(" ", "_")
    out_dir = os.path.join(base_dir, "acciones de empleado", month_folder, emp_nombre)
    os.makedirs(out_dir, exist_ok=True)
    filename = f"{today.strftime('%Y-%m-%d')}_{tipo_doc}_{safe_name}.docx"
    return os.path.join(out_dir, filename)


def _remove_borders(cell):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = tc.makeelement(qn('w:tcBorders'), {})
        tcPr.append(tcBorders)
    for bn in ['top', 'left', 'bottom', 'right']:
        be = tc.makeelement(qn(f'w:{bn}'), {
            qn('w:val'): 'none', qn('w:sz'): '0',
            qn('w:space'): '0', qn('w:color'): 'auto'
        })
        tcBorders.append(be)


def _set_cell_bg(cell, color_hex: str):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    shading = tc.makeelement(qn('w:shd'), {
        qn('w:fill'): color_hex, qn('w:val'): 'clear'
    })
    tcPr.append(shading)


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    margins = tc.makeelement(qn('w:tcMar'), {})
    for side, val in [('top', top), ('bottom', bottom), ('start', left), ('end', right)]:
        m = tc.makeelement(qn(f'w:{side}'), {qn('w:w'): str(val), qn('w:type'): 'dxa'})
        margins.append(m)
    tcPr.append(margins)


def _set_cell_vertical_center(cell):
    tc = cell._element
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = tc.makeelement(qn('w:tcPr'), {})
        tc.insert(0, tcPr)
    vAlign = tc.makeelement(qn('w:vAlign'), {qn('w:val'): 'center'})
    tcPr.append(vAlign)


def _add_thin_line(doc):
    """Adds a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = p._element.makeelement(qn('w:pBdr'), {})
    bottom = p._element.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '4',
        qn('w:space'): '1', qn('w:color'): 'BDC5D6'
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


# ─── Corporate Header & Footer (unified for ALL documents) ───────────────────

def _setup_document(doc: Document, logo_path: str, rrhh_code: str = "", reference_code: str = ""):
    """Sets up page margins, default font, header and footer for the document."""
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    # Global font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)

    # ── HEADER ──
    header = section.header
    header.is_linked_to_previous = False

    # Header table: Logo | Company info (2 columns, cleaner design)
    h_tbl = header.add_table(1, 2, width=Inches(6.5))
    h_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Logo cell (left column)
    logo_cell = h_tbl.cell(0, 0)
    logo_cell.width = Inches(1.8)
    p_logo = logo_cell.paragraphs[0]
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if os.path.exists(logo_path):
        p_logo.add_run().add_picture(logo_path, width=Inches(1.6))
    else:
        r = p_logo.add_run(EMPRESA_NOMBRE)
        r.font.size = Pt(12)
        r.font.bold = True
        r.font.color.rgb = BLUE_DARK

    # Company info cell (right column, stacked vertically)
    info_cell = h_tbl.cell(0, 1)
    info_cell.width = Inches(4.7)
    
    # Reference code (top right)
    if reference_code:
        p_ref = info_cell.paragraphs[0]
        p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r_ref = p_ref.add_run(reference_code)
        r_ref.font.size = Pt(8)
        r_ref.font.bold = True
        r_ref.font.color.rgb = BLUE_MED
    
    # Company name
    p_name = info_cell.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_name = p_name.add_run(EMPRESA_NOMBRE)
    r_name.font.size = Pt(11)
    r_name.font.bold = True
    r_name.font.color.rgb = BLUE_DARK
    
    # Cedula
    p_ced = info_cell.add_paragraph()
    p_ced.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_ced = p_ced.add_run(EMPRESA_CEDULA)
    r_ced.font.size = Pt(8)
    r_ced.font.color.rgb = GRAY_TEXT
    
    # Tel + Address
    p_contact = info_cell.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_contact = p_contact.add_run(EMPRESA_TEL + '  |  ' + EMPRESA_DIR)
    r_contact.font.size = Pt(8)
    r_contact.font.color.rgb = GRAY_TEXT
    
    # RRHH code (if provided)
    if rrhh_code:
        p_code = info_cell.add_paragraph()
        p_code.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rc = p_code.add_run(rrhh_code)
        rc.font.size = Pt(7)
        rc.font.color.rgb = GRAY_TEXT
        rc.font.italic = True


    # Remove borders from header table
    for row in h_tbl.rows:
        for cell in row.cells:
            _remove_borders(cell)
            _set_cell_vertical_center(cell)

    # Header bottom line
    h_line = header.add_paragraph()
    h_line.paragraph_format.space_before = Pt(2)
    h_line.paragraph_format.space_after = Pt(0)
    hPr = h_line._element.get_or_add_pPr()
    hBdr = h_line._element.makeelement(qn('w:pBdr'), {})
    hBottom = h_line._element.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single', qn('w:sz'): '8',
        qn('w:space'): '1', qn('w:color'): '1A3C6E'
    })
    hBdr.append(hBottom)
    hPr.append(hBdr)

    # ── FOOTER ──
    footer = section.footer
    footer.is_linked_to_previous = False

    # Footer top line
    f_line = footer.add_paragraph()
    f_line.paragraph_format.space_before = Pt(0)
    f_line.paragraph_format.space_after = Pt(4)
    fPr = f_line._element.get_or_add_pPr()
    fBdr = f_line._element.makeelement(qn('w:pBdr'), {})
    fTop = f_line._element.makeelement(qn('w:top'), {
        qn('w:val'): 'single', qn('w:sz'): '6',
        qn('w:space'): '1', qn('w:color'): '1A3C6E'
    })
    fBdr.append(fTop)
    fPr.append(fBdr)

    # Footer table: Address | Confidential | Page
    f_tbl = footer.add_table(1, 3, width=Inches(6.5))
    f_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    addr_cell = f_tbl.cell(0, 0)
    p_addr = addr_cell.paragraphs[0]
    p_addr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ra = p_addr.add_run(f"{EMPRESA_DIR}  •  {EMPRESA_TEL}")
    ra.font.size = Pt(7)
    ra.font.color.rgb = GRAY_TEXT

    conf_cell = f_tbl.cell(0, 1)
    p_conf = conf_cell.paragraphs[0]
    p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rc = p_conf.add_run("Documento confidencial — Uso interno")
    rc.font.size = Pt(7)
    rc.font.color.rgb = GRAY_TEXT
    rc.font.italic = True

    page_cell = f_tbl.cell(0, 2)
    p_page = page_cell.paragraphs[0]
    p_page.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp = p_page.add_run("Pág. ")
    rp.font.size = Pt(7)
    rp.font.color.rgb = GRAY_TEXT
    # Page number field
    fldChar1 = p_page._element.makeelement(qn('w:r'), {})
    fldChar1_run = p_page._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'begin'})
    fldChar1.append(fldChar1_run)
    p_page._element.append(fldChar1)
    instrText = p_page._element.makeelement(qn('w:r'), {})
    instrElem = p_page._element.makeelement(qn('w:instrText'), {qn('xml:space'): 'preserve'})
    instrElem.text = ' PAGE '
    instrText.append(instrElem)
    p_page._element.append(instrText)
    fldChar2 = p_page._element.makeelement(qn('w:r'), {})
    fldChar2_run = p_page._element.makeelement(qn('w:fldChar'), {qn('w:fldCharType'): 'end'})
    fldChar2.append(fldChar2_run)
    p_page._element.append(fldChar2)

    for row in f_tbl.rows:
        for cell in row.cells:
            _remove_borders(cell)

    return doc


# ─── Shared building blocks ──────────────────────────────────────────────────

def _add_document_title(doc: Document, title: str):
    tbl = doc.add_table(1, 1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, '1A3C6E')
    _set_cell_vertical_center(cell)
    _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = WHITE


def _add_reference_in_body(doc: Document, reference_code: str):
    """Agrega el código de referencia debajo del título del documento."""
    p_ref = doc.add_paragraph()
    p_ref.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_label = p_ref.add_run("Referencia: ")
    r_label.font.size = Pt(8)
    r_label.font.color.rgb = GRAY_TEXT
    r_code = p_ref.add_run(reference_code)
    r_code.font.size = Pt(9)
    r_code.font.bold = True
    r_code.font.color.rgb = BLUE_MED


def _add_section_header(doc: Document, section_num: str, title: str):
    doc.add_paragraph("")
    tbl = doc.add_table(1, 1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, '2F5496')
    _set_cell_margins(cell, top=20, bottom=20, left=100, right=80)
    p = cell.paragraphs[0]
    run = p.add_run(f"  {section_num}. {title}")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = WHITE


def _add_info_row(doc: Document, label: str, value: str, bold_value=False):
    p = doc.add_paragraph()
    rl = p.add_run(f"  {label}: ")
    rl.font.size = Pt(9.5)
    rl.font.bold = True
    rl.font.color.rgb = BLACK
    rv = p.add_run(str(value))
    rv.font.size = Pt(9.5)
    if bold_value:
        rv.font.bold = True


def _add_data_table(doc: Document, headers: list, rows: list):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.size = Pt(9)
        r.font.bold = True
        r.font.color.rgb = WHITE
        _set_cell_bg(cell, '2F5496')
        _set_cell_margins(cell, top=30, bottom=30)

    for row_idx, row_data in enumerate(rows):
        bg = BLUE_LIGHT_HEX if row_idx % 2 == 0 else 'FFFFFF'
        for col_idx, val in enumerate(row_data):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            _set_cell_bg(cell, bg)
            _set_cell_margins(cell, top=20, bottom=20)

    return table


def _add_clause(doc: Document, number: str, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rn = p.add_run(f"{number} ")
    rn.font.size = Pt(9)
    rn.font.bold = True
    rn.font.color.rgb = BLUE_MED
    rt = p.add_run(text)
    rt.font.size = Pt(9)


def _add_signature_block(doc: Document, emp_nombre: str, emp_cedula: str = ""):
    """Adds a cleaner 4-row signature block with signatures, names, and identifiers."""
    # 4 rows: signatures | labels | names | cedula/cargo
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: Signature lines
    cell_l = table.cell(0, 0)
    p0 = cell_l.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.add_run("_" * 40).font.size = Pt(9)

    cell_r = table.cell(0, 1)
    pr0 = cell_r.paragraphs[0]
    pr0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr0.add_run("_" * 40).font.size = Pt(9)

    # Row 1: Labels (Firma del Colaborador | Por la Empresa)
    cell_l1 = table.cell(1, 0)
    p1 = cell_l1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("Firma del Colaborador(a)")
    r1.font.size = Pt(8)
    r1.font.bold = True

    cell_r1 = table.cell(1, 1)
    pr1 = cell_r1.paragraphs[0]
    pr1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = pr1.add_run(f"Por {EMPRESA_NOMBRE}")
    r_title.font.size = Pt(8)
    r_title.font.bold = True

    # Row 2: Names
    cell_l2 = table.cell(2, 0)
    p2 = cell_l2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"Nombre: {emp_nombre}")
    r2.font.size = Pt(8)

    cell_r2 = table.cell(2, 1)
    pr2 = cell_r2.paragraphs[0]
    pr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr2.add_run("Nombre: ____________________").font.size = Pt(8)

    # Row 3: Cédula / Cargo
    cell_l3 = table.cell(3, 0)
    p3 = cell_l3.paragraphs[0]
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Cédula: {emp_cedula or 'N/A'}")
    r3.font.size = Pt(8)

    cell_r3 = table.cell(3, 1)
    pr3 = cell_r3.paragraphs[0]
    pr3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr3.add_run("Cargo: ____________________").font.size = Pt(8)

    # Remove all borders
    for row in table.rows:
        for cell in row.cells:
            _remove_borders(cell)

    # Date below the table
    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_date = p_date.add_run(f"Fecha: {datetime.now().strftime('%d/%m/%Y')}")
    r_date.font.size = Pt(8)


def _open_folder(path: str):
    try:
        abs_path = os.path.abspath(path)
        subprocess.Popen(f'explorer /select,"{abs_path}"')
    except Exception:
        pass


def _add_page_break(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)


# ═══════════════════════════════════════════════════════════════════════════════
# 1. CONTRATO DE PRÉSTAMO
# ═══════════════════════════════════════════════════════════════════════════════

def generar_prestamo(emp_nombre: str, emp_cedula: str, monto_total: float,
                     pago_semanal: float, logo_path: str, base_dir: str,
                     rrhh_code: str = "") -> str:
    doc = Document()
    ref_code = get_next_reference("prestamo")
    _setup_document(doc, logo_path, rrhh_code, ref_code)
    _add_reference_in_body(doc, ref_code)

    semanas = math.ceil(monto_total / pago_semanal) if pago_semanal > 0 else 0
    hoy = datetime.now()
    fecha_liquidacion = (hoy + timedelta(weeks=semanas)).strftime("%d/%m/%Y")

    _add_document_title(doc, "PRÉSTAMO A EMPLEADO")

    _add_section_header(doc, "I", "DATOS DEL COLABORADOR")
    _add_info_row(doc, "Nombre completo", emp_nombre, bold_value=True)
    _add_info_row(doc, "Cédula de identidad", emp_cedula or "N/A")
    _add_info_row(doc, "Fecha del documento", hoy.strftime("%d/%m/%Y"))

    _add_section_header(doc, "II", "DATOS DEL PRÉSTAMO")
    _add_info_row(doc, "Monto total del préstamo", f"₡{monto_total:,.2f}", bold_value=True)
    _add_info_row(doc, "Fecha de desembolso", hoy.strftime("%d/%m/%Y"))
    _add_info_row(doc, "Forma de desembolso", "Deducción directa de planilla semanal")

    _add_section_header(doc, "III", "PLAN DE PAGO SEMANAL")
    _add_info_row(doc, "Monto de cuota semanal", f"₡{pago_semanal:,.2f}", bold_value=True)
    _add_info_row(doc, "Primera cuota estimada", (hoy + timedelta(weeks=1)).strftime("%d/%m/%Y"))
    _add_info_row(doc, "Última cuota estimada", fecha_liquidacion)
    _add_info_row(doc, "Número total de cuotas", str(semanas))

    _add_section_header(doc, "IV", "CONDICIONES Y CLÁUSULAS DEL PRÉSTAMO")
    _add_clause(doc, "4.1",
        f"La empresa concede al colaborador el presente préstamo de manera excepcional, de conformidad "
        f"con su política interna y en atención a la solicitud formal del colaborador.")
    _add_clause(doc, "4.2",
        f"El monto total adeudado será deducido íntegramente de la planilla semanal del colaborador, en "
        f"cuotas fijas de ₡{pago_semanal:,.2f}. Se entiende que el colaborador autoriza expresamente dichas "
        f"deducciones a su salario base.")
    _add_clause(doc, "4.3",
        f"En caso de terminación de la relación laboral, por cualquier causa, el saldo pendiente del préstamo "
        f"será deducido en su totalidad de la liquidación final de los derechos laborales correspondientes, "
        f"incluyendo salario, vacaciones, aguinaldo y demás prestaciones del colaborador.")
    _add_clause(doc, "4.4",
        f"En un plazo máximo de dos (2) meses consecutivos sin cancelar las cuotas parciales, sea por imprevisto, "
        f"ausencias injustificadas o cualquier otra razón, la empresa podrá exigir el pago total inmediato del "
        f"saldo restante conforme a la ley.")
    _add_clause(doc, "4.5",
        f"Este documento constituye pleno reconocimiento de la deuda por parte del colaborador y tendrá plena "
        f"validez legal como título ejecutivo.")
    _add_clause(doc, "4.6",
        f"El colaborador podrá realizar pagos extraordinarios o de cancelación anticipada, sin penalidad, "
        f"notificando por escrito al Departamento de Recursos Humanos con al menos tres (3) días hábiles de anticipación.")
    _add_clause(doc, "4.7",
        f"El colaborador autoriza expresamente a {EMPRESA_NOMBRE} para realizar las deducciones de planilla que "
        f"correspondan según el plan de pago establecido en el presente documento.")

    _add_section_header(doc, "V", "DECLARACIÓN DEL COLABORADOR")
    doc.add_paragraph("")
    p_decl = doc.add_paragraph()
    p_decl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rd = p_decl.add_run(
        f"El(la) colaborador(a), declara bajo fe de juramento que la información aquí suministrada es veraz y exacta, "
        f"que ha leído y comprendido todos los términos y condiciones del presente contrato, y que autoriza las deducciones "
        f"de planilla estipuladas en la Sección III y las cláusulas del apartado IV anterior. Asimismo, declara haber "
        f"recibido copia íntegra del presente documento a su entera satisfacción.")
    rd.font.size = Pt(9)

    doc.add_paragraph("")
    p_lf = doc.add_paragraph()
    rlf = p_lf.add_run("Lugar y fecha: ________________________________________")
    rlf.font.size = Pt(9)

    _add_signature_block(doc, emp_nombre, emp_cedula)

    # ── HOJA APARTE: TABLA DE AMORTIZACIÓN ──
    doc.add_page_break()
    _add_document_title(doc, "TABLA DE AMORTIZACIÓN DE PAGOS")
    doc.add_paragraph("")
    _add_info_row(doc, "Colaborador", emp_nombre, bold_value=True)
    _add_info_row(doc, "Monto del préstamo", f"₡{monto_total:,.2f}")
    _add_info_row(doc, "Cuota semanal", f"₡{pago_semanal:,.2f}")
    doc.add_paragraph("")

    amort_rows = []
    saldo = monto_total
    for i in range(1, semanas + 1):
        fecha_cuota = (hoy + timedelta(weeks=i)).strftime("%d/%m/%Y")
        cuota = min(pago_semanal, saldo)
        abonado = monto_total - saldo + cuota
        saldo -= cuota
        amort_rows.append([
            str(i), fecha_cuota, f"₡{cuota:,.2f}",
            f"₡{abonado:,.2f}", f"₡{max(saldo, 0):,.2f}", ""
        ])

    _add_data_table(doc,
        ["#", "Fecha de pago", "Cuota (₡)", "Abonado (₡)", "Saldo (₡)", "Firma"],
        amort_rows)

    out_path = _build_output_path(base_dir, "Prestamo", emp_nombre)
    doc.save(out_path)
    _open_folder(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# 2. CARTA DE AMONESTACIÓN POR FALTANTES / TARDANZAS / CONDUCTAS
# ═══════════════════════════════════════════════════════════════════════════════

def generar_amonestacion(emp_nombre: str, emp_cedula: str, tipo: str,
                         datos: list, logo_path: str, base_dir: str,
                         rrhh_code: str = "") -> str:
    doc = Document()
    ref_code = get_next_reference("amonestacion")
    _setup_document(doc, logo_path, rrhh_code, ref_code)
    _add_reference_in_body(doc, ref_code)

    hoy = datetime.now()

    titles = {
        "faltantes": "CARTA DE AMONESTACIÓN POR FALTANTES",
        "tardanzas": "CARTA DE AMONESTACIÓN POR LLEGADAS TARDÍAS",
        "conductas": "CARTA POR CONDUCTAS INAPROPIADAS"
    }
    title = titles.get(tipo, "CARTA DE AMONESTACIÓN")
    _add_document_title(doc, title)

    _add_section_header(doc, "I", "DATOS DEL COLABORADOR")
    _add_info_row(doc, "Nombre completo", emp_nombre, bold_value=True)
    _add_info_row(doc, "Cédula de identidad", emp_cedula or "N/A")
    _add_info_row(doc, "Fecha de emisión", hoy.strftime("%d/%m/%Y"))

    _add_section_header(doc, "II", "ANTECEDENTES Y HECHOS")
    doc.add_paragraph("")
    p_ant = doc.add_paragraph()
    p_ant.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    if tipo == "faltantes":
        r_ant = p_ant.add_run(
            f"La administración de {EMPRESA_NOMBRE} ha detectado, a través de los procedimientos de arqueo, "
            f"auditoría y control interno, una serie de diferencias negativas (faltantes) en los valores, "
            f"inventarios o caja bajo la responsabilidad directa del colaborador ")
        r_ant.font.size = Pt(9)
        r_name = p_ant.add_run(f"{emp_nombre}")
        r_name.font.size = Pt(9)
        r_name.font.bold = True
        r_cont = p_ant.add_run(". Los hechos se documentan en el siguiente detalle:")
        r_cont.font.size = Pt(9)

        _add_section_header(doc, "III", "DETALLE DE FALTANTES REGISTRADOS")
        doc.add_paragraph("")
        total = sum(float(f.get("monto", 0)) for f in datos)
        rows = []
        for i, f in enumerate(datos):
            monto = float(f.get("monto", 0))
            rows.append([str(i + 1), str(f.get("fecha", "")), f"₡{monto:,.2f}"])
        rows.append(["", "TOTAL GENERAL", f"₡{total:,.2f}"])
        _add_data_table(doc, ["#", "Fecha del Faltante", "Monto (₡)"], rows)

    elif tipo == "tardanzas":
        r_ant = p_ant.add_run(
            f"El Departamento de Recursos Humanos de {EMPRESA_NOMBRE} ha detectado, mediante el sistema "
            f"de control de asistencia, una serie de llegadas tardías injustificadas por parte del colaborador ")
        r_ant.font.size = Pt(9)
        r_name = p_ant.add_run(f"{emp_nombre}")
        r_name.font.size = Pt(9)
        r_name.font.bold = True
        r_cont = p_ant.add_run(". El registro de tardanzas es el siguiente:")
        r_cont.font.size = Pt(9)

        _add_section_header(doc, "III", "REGISTRO DE LLEGADAS TARDÍAS")
        doc.add_paragraph("")
        total_mins = sum(int(f.get("minutos", 0)) for f in datos)
        rows = []
        for i, f in enumerate(datos):
            mins = int(f.get("minutos", 0))
            rows.append([str(i + 1), str(f.get("fecha", "")), f"{mins} min"])
        rows.append(["", "TOTAL ACUMULADO", f"{total_mins} min"])
        _add_data_table(doc, ["#", "Fecha de Tardanza", "Minutos Tarde"], rows)

    elif tipo == "conductas":
        r_ant = p_ant.add_run(
            f"Por medio del presente documento, la Gerencia de {EMPRESA_NOMBRE} procede a notificar de "
            f"manera formal al colaborador ")
        r_ant.font.size = Pt(9)
        r_name = p_ant.add_run(f"{emp_nombre}")
        r_name.font.size = Pt(9)
        r_name.font.bold = True
        r_cont = p_ant.add_run(" sobre conductas recientes que se consideran inapropiadas o perjudiciales para "
                               "el buen funcionamiento y clima laboral de la empresa.")
        r_cont.font.size = Pt(9)

        _add_section_header(doc, "III", "DESCRIPCIÓN DE LA SITUACIÓN O CONDUCTA INAPROPIADA")
        doc.add_paragraph("")
        p_desc = doc.add_paragraph()
        p_desc.add_run("(Espacio reservado para redacción manual por parte de la Jefatura)").font.color.rgb = GRAY_TEXT
        for _ in range(8):
            p_line = doc.add_paragraph()
            p_line.add_run("_" * 80).font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)

    _add_section_header(doc, "IV", "FUNDAMENTO Y DISPOSICIONES DISCIPLINARIAS")
    if tipo == "faltantes":
        _add_clause(doc, "4.1", "La custodia, manejo y rendición de cuentas de los valores asignados al colaborador constituyen obligaciones esenciales derivadas de su contrato de trabajo.")
        _add_clause(doc, "4.2", "El monto total de los faltantes indicados deberá ser reintegrado a la empresa conforme al acuerdo que se establezca entre ambas partes.")
    elif tipo == "tardanzas":
        _add_clause(doc, "4.1", "El cumplimiento estricto del horario de trabajo asignado es una obligación laboral ineludible según lo pactado en su contrato y el reglamento interno.")
        _add_clause(doc, "4.2", "Las llegadas tardías constantes afectan significativamente la operatividad del equipo de trabajo y el servicio prestado por la empresa.")
    elif tipo == "conductas":
        _add_clause(doc, "4.1", "Las normas de sana convivencia, el respeto hacia los compañeros, superiores y clientes, así como el cumplimiento estricto de los reglamentos internos son pilares fundamentales de la cultura organizacional.")
        _add_clause(doc, "4.2", "Cualquier comportamiento que atente contra estos principios representa un incumplimiento severo de las disposiciones laborales vigentes.")

    _add_clause(doc, "4.3", "La presente misiva se emite como una medida preventiva formal y quedará registrada en el expediente personal del colaborador.")
    _add_clause(doc, "4.4", "Se le exhorta formalmente a tomar las medidas correctivas inmediatas para evitar la reincidencia. La reiteración de este comportamiento podría facultar a la empresa a tomar acciones disciplinarias de mayor severidad, hasta llegar al despido sin responsabilidad patronal si los hechos así lo ameritan conforme a la ley.")

    _add_section_header(doc, "V", "ACUSE DE RECIBO DEL COLABORADOR")
    doc.add_paragraph("")
    p_acuse = doc.add_paragraph()
    p_acuse.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_acuse = p_acuse.add_run(
        f"Yo, {emp_nombre}, con cédula de identidad {emp_cedula or 'N/A'}, confirmo haber sido debidamente "
        f"notificado(a) sobre los hechos descritos en la presente carta. Comprendo que este "
        f"documento será incluido en mi expediente laboral y que la empresa se reserva las acciones correspondientes "
        f"en caso de reincidencia.")
    r_acuse.font.size = Pt(9)

    doc.add_paragraph("")
    p_lf = doc.add_paragraph()
    rlf = p_lf.add_run("Lugar y fecha: ________________________________________")
    rlf.font.size = Pt(9)

    _add_signature_block(doc, emp_nombre, emp_cedula)

    out_path = _build_output_path(base_dir, "Amonestacion", emp_nombre)
    doc.save(out_path)
    _open_folder(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# 3. CONSTANCIA DE VACACIONES
# ═══════════════════════════════════════════════════════════════════════════════

def generar_vacaciones(emp_nombre: str, emp_cedula: str, tipo: str,
                       fecha_inicio: str, fecha_reingreso: str,
                       logo_path: str, base_dir: str,
                       rrhh_code: str = "") -> str:
    doc = Document()
    ref_code = get_next_reference("vacaciones")
    _setup_document(doc, logo_path, rrhh_code, ref_code)
    _add_reference_in_body(doc, ref_code)

    hoy = datetime.now()
    try:
        dt_inicio = datetime.strptime(fecha_inicio, "%Y-%m-%d")
        dt_reingreso = datetime.strptime(fecha_reingreso, "%Y-%m-%d")
        dias = (dt_reingreso - dt_inicio).days
    except Exception:
        dias = 0
        dt_inicio = hoy
        dt_reingreso = hoy

    tipo_label = "Goce Salarial Total" if tipo == "total" else "Goce Salarial Parcial / Fraccionado"

    _add_document_title(doc, "CONSTANCIA DE VACACIONES")

    _add_section_header(doc, "I", "DATOS DEL COLABORADOR")
    _add_info_row(doc, "Nombre completo", emp_nombre, bold_value=True)
    _add_info_row(doc, "Cédula de identidad", emp_cedula or "N/A")
    _add_info_row(doc, "Fecha de emisión", hoy.strftime("%d/%m/%Y"))

    _add_section_header(doc, "II", "DETALLE DEL PERÍODO VACACIONAL")
    doc.add_paragraph("")
    _add_data_table(doc,
        ["Concepto", "Descripción"],
        [
            ["Tipo de aplicación", tipo_label],
            ["Fecha de inicio de vacaciones", dt_inicio.strftime("%d/%m/%Y")],
            ["Fecha de reingreso laboral", dt_reingreso.strftime("%d/%m/%Y")],
            ["Total de días naturales", f"{dias} días"],
        ])

    _add_section_header(doc, "III", "CONDICIONES Y OBSERVACIONES")
    _add_clause(doc, "3.1",
        f"El presente documento certifica que el colaborador {emp_nombre} ha coordinado con la "
        f"administración de {EMPRESA_NOMBRE} el disfrute de su período anual de vacaciones, en cumplimiento "
        f"de lo establecido por el Código de Trabajo de Costa Rica y las políticas internas de la empresa.")
    _add_clause(doc, "3.2",
        f"El colaborador cesará temporalmente la prestación de sus servicios a partir del "
        f"{dt_inicio.strftime('%d/%m/%Y')}, debiendo reincorporarse a su jornada ordinaria el "
        f"{dt_reingreso.strftime('%d/%m/%Y')} en el horario que le corresponda según su turno asignado.")
    _add_clause(doc, "3.3",
        f"Los {dias} día(s) de descanso aquí consignados serán rebajados del saldo total de días de "
        f"vacaciones acumulados por el colaborador. El tipo de aplicación es: \"{tipo_label}\".")
    _add_clause(doc, "3.4",
        f"El colaborador declara estar conforme con las fechas aprobadas y se compromete a reintegrarse "
        f"puntualmente en la fecha indicada. La inasistencia posterior a la fecha de reingreso sin justificación "
        f"será procesada conforme a la normativa disciplinaria vigente.")

    _add_section_header(doc, "IV", "DECLARACIÓN DE CONFORMIDAD")
    doc.add_paragraph("")
    p_decl = doc.add_paragraph()
    p_decl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_decl = p_decl.add_run(
        f"Ambas partes — el colaborador y la representación legal de {EMPRESA_NOMBRE} — manifiestan su plena "
        f"conformidad con las condiciones del presente documento. Se extiende la presente para los efectos legales "
        f"y administrativos correspondientes, y para ser integrada al expediente laboral del colaborador.")
    r_decl.font.size = Pt(9)

    doc.add_paragraph("")
    p_lf = doc.add_paragraph()
    rlf = p_lf.add_run("Lugar y fecha: ________________________________________")
    rlf.font.size = Pt(9)

    _add_signature_block(doc, emp_nombre, emp_cedula)

    out_path = _build_output_path(base_dir, "Vacaciones", emp_nombre)
    doc.save(out_path)
    _open_folder(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# 4/5. CARTA DE DESPIDO / RENUNCIA
# ═══════════════════════════════════════════════════════════════════════════════

def generar_liquidacion(tipo_doc: str, emp_nombre: str, emp_cedula: str,
                        vac_dias: float, vac_monto: float, aguinaldo_monto: float,
                        cesantia_monto: float, preaviso_monto: float,
                        total_pagar: float, modo_pago: str,
                        logo_path: str, base_dir: str,
                        rrhh_code: str = "") -> str:
    doc = Document()
    ref_code = get_next_reference("despido" if tipo_doc == "Despido" else "renuncia")
    _setup_document(doc, logo_path, rrhh_code, ref_code)
    _add_reference_in_body(doc, ref_code)

    hoy = datetime.now()

    if tipo_doc == "Despido":
        titulo_principal = "CARTA DE DESPIDO"
        texto_intro = f"Por medio del presente documento, {EMPRESA_NOMBRE} le notifica formalmente su despido y la terminación de su contrato de trabajo."
    else:
        titulo_principal = "CARTA DE ACEPTACIÓN DE RENUNCIA"
        texto_intro = f"Por medio del presente documento, {EMPRESA_NOMBRE} acusa recibo y acepta formalmente su renuncia voluntaria al puesto que desempeñaba."

    _add_document_title(doc, titulo_principal)

    _add_section_header(doc, "I", "DATOS DEL COLABORADOR")
    _add_info_row(doc, "Nombre completo", emp_nombre, bold_value=True)
    _add_info_row(doc, "Cédula de identidad", emp_cedula or "N/A")
    _add_info_row(doc, "Fecha de emisión", hoy.strftime("%d/%m/%Y"))

    _add_section_header(doc, "II", "DETALLE DE TERMINACIÓN LABORAL")
    doc.add_paragraph("")
    p_ant = doc.add_paragraph()
    p_ant.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_ant = p_ant.add_run(texto_intro + f" La acción se hará efectiva a partir del día de percibo de este documento. Agradecemos los servicios prestados durante su permanencia en la empresa.")
    r_ant.font.size = Pt(9)

    _add_section_header(doc, "III", "DETALLE DE DERECHOS Y LIQUIDACIÓN")
    doc.add_paragraph("")
    rows = [
        ["Vacaciones Proporcionales", f"{vac_dias:.2f} días", f"₡{vac_monto:,.2f}"],
        ["Aguinaldo Proporcional", "Monto acumulado", f"₡{aguinaldo_monto:,.2f}"],
    ]
    if tipo_doc == "Despido":
        if cesantia_monto > 0:
            rows.append(["Auxilio de Cesantía (Art. 29 CT)", "Según antigüedad", f"₡{cesantia_monto:,.2f}"])
        if preaviso_monto > 0:
            rows.append(["Preaviso", "Según antigüedad", f"₡{preaviso_monto:,.2f}"])
    rows.append(["", "TOTAL A PAGAR", f"₡{total_pagar:,.2f}"])
    _add_data_table(doc, ["Concepto", "Detalle", "Monto Calculado (₡)"], rows)

    _add_section_header(doc, "IV", "CONDICIONES DE PAGO")
    _add_clause(doc, "4.1", f"El monto total estipulado será cancelado bajo la modalidad de: {modo_pago}.")
    _add_clause(doc, "4.2", "Con este pago, la empresa cancela de manera definitiva y plena todos los extremos laborales adquiridos durante el período laborado por el empleado, no adeudándosele suma alguna por concepto de cesantía, preaviso o cualquier otro adicional.")
    _add_clause(doc, "4.3", "Al momento del retiro, el colaborador deberá hacer entrega formal de cualquier herramienta, indumentaria o equipo de trabajo provisto por la empresa.")

    _add_section_header(doc, "V", "CONFORMIDAD")
    doc.add_paragraph("")
    p_decl = doc.add_paragraph()
    p_decl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_decl = p_decl.add_run(
        f"Yo, {emp_nombre}, declaro estar cien por ciento conforme con las condiciones del presente documento "
        f"así como con los rubros calculados en mi liquidación. Al recibir dichos montos, no tendré nada más que reclamar "
        f"en la vía administrativa, judicial o laboral a {EMPRESA_NOMBRE}.")
    r_decl.font.size = Pt(9)

    doc.add_paragraph("")
    p_lf = doc.add_paragraph()
    rlf = p_lf.add_run("Lugar y fecha: ________________________________________")
    rlf.font.size = Pt(9)

    _add_signature_block(doc, emp_nombre, emp_cedula)

    safe_title = "Despido" if tipo_doc == "Despido" else "Renuncia"
    out_path = _build_output_path(base_dir, safe_title, emp_nombre)
    doc.save(out_path)
    _open_folder(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# 6. CARTA DE RECOMENDACIÓN
# ═══════════════════════════════════════════════════════════════════════════════

def generar_recomendacion(emp_nombre: str, emp_cedula: str, puesto: str,
                          fecha_inicio: str, texto_adicional: str,
                          logo_path: str, base_dir: str,
                          rrhh_code: str = ""):
    doc = Document()
    ref_code = get_next_reference("recomendacion")
    _setup_document(doc, logo_path, rrhh_code, ref_code)
    _add_reference_in_body(doc, ref_code)

    hoy = datetime.now()
    meses_es = {
        "January": "enero", "February": "febrero", "March": "marzo",
        "April": "abril", "May": "mayo", "June": "junio",
        "July": "julio", "August": "agosto", "September": "septiembre",
        "October": "octubre", "November": "noviembre", "December": "diciembre"
    }
    fecha_doc = hoy.strftime("%d de %B de %Y")
    for en, es in meses_es.items():
        fecha_doc = fecha_doc.replace(en, es)

    _add_document_title(doc, "CARTA DE RECOMENDACIÓN")

    p_date = doc.add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_date = p_date.add_run(f"San Carlos, {fecha_doc}")
    r_date.font.size = Pt(10)

    doc.add_paragraph()
    p_to = doc.add_paragraph()
    p_to.add_run("A QUIEN INTERESE").bold = True
    p_to.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph()

    antiguedad_text = ""
    if fecha_inicio:
        try:
            fi = datetime.strptime(fecha_inicio, "%Y-%m-%d")
            diff = hoy - fi
            anios = diff.days // 365
            meses = (diff.days % 365) // 30
            if anios > 0:
                antiguedad_text = f"{anios} año(s) y {meses} mes(es)"
            else:
                antiguedad_text = f"{meses} mes(es)"
        except Exception:
            antiguedad_text = ""

    body_text = (
        f"Por medio de la presente, Servicentro La Marina, cédula jurídica 3-101-130178, "
        f"hace constar que el(la) señor(a) {emp_nombre}, portador(a) de la cédula de identidad "
        f"número {emp_cedula}, laboró en nuestra empresa desempeñando el puesto de "
        f"{puesto}")
    if fecha_inicio and antiguedad_text:
        body_text += f", desde el {fecha_inicio} hasta la fecha, acumulando una antigüedad de {antiguedad_text}"
    body_text += "."

    p_body = doc.add_paragraph()
    p_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_body = p_body.add_run(body_text)
    r_body.font.size = Pt(10)

    doc.add_paragraph()
    rec_text = (
        f"Durante su período laboral, el(la) señor(a) {emp_nombre} demostró ser una persona "
        f"responsable, honesta, puntual y dedicada en el cumplimiento de sus funciones. "
        f"Su desempeño fue satisfactorio y mantuvo excelentes relaciones interpersonales "
        f"con sus compañeros de trabajo y superiores.")
    p_rec = doc.add_paragraph()
    p_rec.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_rec = p_rec.add_run(rec_text)
    r_rec.font.size = Pt(10)

    if texto_adicional:
        doc.add_paragraph()
        p_extra = doc.add_paragraph()
        p_extra.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r_extra = p_extra.add_run(texto_adicional)
        r_extra.font.size = Pt(10)

    doc.add_paragraph()
    p_close = doc.add_paragraph()
    p_close.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_close = p_close.add_run(
        "Se extiende la presente carta de recomendación a solicitud del interesado(a), "
        "para los fines que estime convenientes.")
    r_close.font.size = Pt(10)

    doc.add_paragraph()
    doc.add_paragraph()

    _add_signature_block(doc, emp_nombre, emp_cedula)

    out_path = _build_output_path(base_dir, "Recomendacion", emp_nombre)
    doc.save(out_path)
    _open_folder(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# 7. LIQUIDACIÓN LABORAL (documento independiente)
# ═══════════════════════════════════════════════════════════════════════════════

def generar_liquidacion_laboral(emp_nombre: str, emp_cedula: str,
                                cesantia_monto: float,
                                vac_monto: float, vac_dias: float,
                                aguinaldo_monto: float,
                                preaviso_monto: float,
                                logo_path: str, base_dir: str,
                                rrhh_code: str = "") -> str:
    """Genera carta de liquidación laboral independiente.

    - vac_dias: si es 0 o None, no se mencionan los días
    - preaviso_monto: si es 0 o None, no se incluye la fila
    """
    doc = Document()
    ref_code = get_next_reference("liquidacion_laboral")
    _setup_document(doc, logo_path, rrhh_code, ref_code)
    _add_reference_in_body(doc, ref_code)

    hoy = datetime.now()

    _add_document_title(doc, "LIQUIDACIÓN LABORAL")

    _add_section_header(doc, "I", "DATOS DEL COLABORADOR")
    _add_info_row(doc, "Nombre completo", emp_nombre, bold_value=True)
    _add_info_row(doc, "Cédula de identidad", emp_cedula or "N/A")
    _add_info_row(doc, "Fecha de emisión", hoy.strftime("%d/%m/%Y"))

    _add_section_header(doc, "II", "DETALLE DE LIQUIDACIÓN")
    doc.add_paragraph("")

    total = cesantia_monto + vac_monto + aguinaldo_monto + (preaviso_monto or 0)

    rows = []
    # Cesantía
    rows.append(["Auxilio de Cesantía", "Según antigüedad", f"₡{cesantia_monto:,.2f}"])

    # Vacaciones
    if vac_dias and vac_dias > 0:
        rows.append(["Vacaciones", f"{vac_dias:.1f} días", f"₡{vac_monto:,.2f}"])
    else:
        rows.append(["Vacaciones", "Monto acumulado", f"₡{vac_monto:,.2f}"])

    # Aguinaldo proporcional
    rows.append(["Aguinaldo Proporcional", "Según período", f"₡{aguinaldo_monto:,.2f}"])

    # Preaviso (solo si aplica)
    if preaviso_monto and preaviso_monto > 0:
        rows.append(["Preaviso", "Según antigüedad", f"₡{preaviso_monto:,.2f}"])

    rows.append(["", "TOTAL LIQUIDACIÓN", f"₡{total:,.2f}"])

    _add_data_table(doc, ["Concepto", "Detalle", "Monto (₡)"], rows)

    _add_section_header(doc, "III", "OBSERVACIONES LEGALES")
    _add_clause(doc, "3.1",
        "El presente documento detalla los montos correspondientes a la liquidación de derechos laborales "
        "del colaborador, conforme a lo estipulado en el Código de Trabajo de Costa Rica.")
    _add_clause(doc, "3.2",
        f"Los montos aquí indicados constituyen la totalidad de las prestaciones adeudadas por {EMPRESA_NOMBRE} "
        f"al colaborador en concepto de terminación de la relación laboral.")
    _add_clause(doc, "3.3",
        "Cualquier deuda pendiente del colaborador con la empresa (préstamos, faltantes, mercadería de crédito) "
        "será deducida del monto total de la liquidación previo al desembolso final.")

    _add_section_header(doc, "IV", "CONFORMIDAD Y ACEPTACIÓN")
    doc.add_paragraph("")
    p_decl = doc.add_paragraph()
    p_decl.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_decl = p_decl.add_run(
        f"Yo, {emp_nombre}, con cédula {emp_cedula or 'N/A'}, declaro haber recibido a mi entera satisfacción "
        f"los montos detallados en la sección II del presente documento. Con la recepción de esta liquidación, "
        f"manifiesto que no tengo nada más que reclamar a {EMPRESA_NOMBRE} por concepto de prestaciones "
        f"laborales, salarios, vacaciones, aguinaldo, cesantía, preaviso ni ningún otro extremo laboral.")
    r_decl.font.size = Pt(9)

    doc.add_paragraph("")
    p_lf = doc.add_paragraph()
    rlf = p_lf.add_run("Lugar y fecha: ________________________________________")
    rlf.font.size = Pt(9)

    _add_signature_block(doc, emp_nombre, emp_cedula)

    out_path = _build_output_path(base_dir, "Liquidacion_Laboral", emp_nombre)
    doc.save(out_path)
    _open_folder(out_path)
    return out_path
